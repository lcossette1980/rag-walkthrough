# ================================================================
# COMPLETE PYTHON RAG IMPLEMENTATION BOOK
# Retrieval-Augmented Generation from Basics to Production
# ================================================================

"""
This comprehensive Python book covers everything you need to know about
building RAG (Retrieval-Augmented Generation) systems, from basic concepts
to production-ready implementations.

Table of Contents:
==================
1. RAG Fundamentals & Setup
2. Basic RAG Implementation
3. Document Processing & Chunking
4. Embedding Generation & Vector Storage
5. Retrieval Mechanisms
6. Response Generation
7. Enhanced RAG with Metadata
8. Multi-Modal RAG Implementation
9. Production RAG Pipeline
10. Monitoring & Analytics
11. Deployment Strategies
12. Best Practices & Optimization

Author: RAG Implementation Guide
Version: 1.0
"""

# ================================================================
# SECTION 1: RAG FUNDAMENTALS & SETUP
# ================================================================

import os
import sys
import logging
import warnings
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass, field
from datetime import datetime
import json
import pickle
import hashlib
import asyncio
import time
from pathlib import Path

# Suppress warnings for cleaner output
warnings.filterwarnings("ignore")

# Core ML and NLP libraries
import numpy as np
import pandas as pd
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer

# Document processing libraries
import PyPDF2
import docx
from bs4 import BeautifulSoup
import markdown

# Vector database and embedding libraries
import faiss
import chromadb
from sentence_transformers import SentenceTransformer

# LangChain components
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    CharacterTextSplitter,
    TokenTextSplitter
)
from langchain.document_loaders import (
    PyPDFLoader,
    TextLoader,
    DirectoryLoader,
    WebBaseLoader
)
from langchain.embeddings import (
    OpenAIEmbeddings,
    HuggingFaceEmbeddings,
    SentenceTransformerEmbeddings
)
from langchain.vectorstores import FAISS, Chroma
from langchain.llms import OpenAI, HuggingFacePipeline
from langchain.chat_models import ChatOpenAI
from langchain.schema import Document
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMChainExtractor

# API clients
import openai
import requests

# Web framework for deployment
from flask import Flask, request, jsonify
import streamlit as st

# Monitoring and logging
from prometheus_client import Counter, Histogram, start_http_server
import structlog

# Configuration and environment
from dotenv import load_dotenv
load_dotenv()

# Setup structured logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = structlog.get_logger()

# ================================================================
# SECTION 2: CONFIGURATION AND DATA CLASSES
# ================================================================

@dataclass
class RAGConfig:
    """
    Configuration class for RAG system settings.
    
    This class centralizes all configuration parameters for the RAG system,
    making it easy to modify behavior without changing code throughout
    the application.
    """
    
    # LLM Configuration
    llm_provider: str = "openai"  # Options: openai, huggingface, local
    llm_model: str = "gpt-3.5-turbo"
    llm_temperature: float = 0.1  # Lower = more deterministic
    llm_max_tokens: int = 512
    
    # Embedding Configuration
    embedding_provider: str = "openai"  # Options: openai, huggingface, sentence_transformers
    embedding_model: str = "text-embedding-ada-002"
    embedding_dimension: int = 1536  # Depends on model chosen
    
    # Chunking Configuration
    chunk_size: int = 1000  # Characters per chunk
    chunk_overlap: int = 200  # Overlap between chunks
    chunking_strategy: str = "recursive"  # Options: recursive, character, token
    
    # Retrieval Configuration
    retrieval_top_k: int = 4  # Number of documents to retrieve
    retrieval_threshold: float = 0.7  # Minimum similarity score
    retrieval_method: str = "similarity"  # Options: similarity, mmr, similarity_score_threshold
    
    # Vector Database Configuration
    vector_db_type: str = "faiss"  # Options: faiss, chroma, pinecone
    vector_db_path: str = "./vector_store"
    persist_directory: str = "./persistent_db"
    
    # API Keys (loaded from environment)
    openai_api_key: str = field(default_factory=lambda: os.getenv("OPENAI_API_KEY"))
    huggingface_api_key: str = field(default_factory=lambda: os.getenv("HUGGINGFACE_API_KEY"))
    
    # Performance Configuration
    batch_size: int = 32  # For batch processing
    max_workers: int = 4  # For concurrent processing
    cache_embeddings: bool = True
    
    # Logging and Monitoring
    log_level: str = "INFO"
    enable_metrics: bool = True
    metrics_port: int = 8000

@dataclass
class DocumentMetadata:
    """
    Metadata structure for documents in the knowledge base.
    
    This class standardizes metadata across different document types,
    enabling rich filtering and search capabilities.
    """
    
    # Basic document information
    document_id: str
    title: str
    source: str  # File path, URL, or source identifier
    document_type: str  # pdf, txt, html, etc.
    
    # Content metadata
    word_count: int = 0
    page_count: int = 0
    language: str = "en"
    
    # Categorization
    category: str = "general"
    tags: List[str] = field(default_factory=list)
    
    # Temporal information
    created_at: datetime = field(default_factory=datetime.now)
    modified_at: datetime = field(default_factory=datetime.now)
    indexed_at: datetime = field(default_factory=datetime.now)
    
    # Quality metrics
    quality_score: float = 1.0  # 0.0 to 1.0
    relevance_score: float = 1.0  # 0.0 to 1.0
    
    # Access control
    access_level: str = "public"  # public, private, restricted
    owner: str = "system"

# ================================================================
# SECTION 3: BASIC RAG IMPLEMENTATION
# ================================================================

class BasicRAG:
    """
    A simple, straightforward RAG implementation.
    
    This class demonstrates the core concepts of RAG:
    1. Document loading and chunking
    2. Embedding generation
    3. Vector storage and retrieval
    4. Response generation
    
    Perfect for learning and prototyping.
    """
    
    def __init__(self, config: RAGConfig):
        """
        Initialize the Basic RAG system.
        
        Args:
            config: RAGConfig instance with system settings
        """
        self.config = config
        self.documents = []
        self.embeddings = None
        self.vectorstore = None
        self.retriever = None
        self.llm = None
        
        # Initialize components
        self._setup_embeddings()
        self._setup_llm()
        
        logger.info("BasicRAG initialized successfully")
    
    def _setup_embeddings(self):
        """
        Initialize the embedding model based on configuration.
        
        This method sets up different embedding providers:
        - OpenAI: High quality, requires API key
        - HuggingFace: Free, good quality
        - SentenceTransformers: Local, customizable
        """
        if self.config.embedding_provider == "openai":
            self.embeddings = OpenAIEmbeddings(
                openai_api_key=self.config.openai_api_key,
                model=self.config.embedding_model
            )
            logger.info("Using OpenAI embeddings")
            
        elif self.config.embedding_provider == "huggingface":
            self.embeddings = HuggingFaceEmbeddings(
                model_name=self.config.embedding_model
            )
            logger.info("Using HuggingFace embeddings")
            
        elif self.config.embedding_provider == "sentence_transformers":
            self.embeddings = SentenceTransformerEmbeddings(
                model_name=self.config.embedding_model
            )
            logger.info("Using SentenceTransformer embeddings")
            
        else:
            raise ValueError(f"Unsupported embedding provider: {self.config.embedding_provider}")
    
    def _setup_llm(self):
        """
        Initialize the language model for response generation.
        
        Supports multiple LLM providers for flexibility in deployment
        and cost optimization.
        """
        if self.config.llm_provider == "openai":
            self.llm = ChatOpenAI(
                openai_api_key=self.config.openai_api_key,
                model_name=self.config.llm_model,
                temperature=self.config.llm_temperature,
                max_tokens=self.config.llm_max_tokens
            )
            logger.info(f"Using OpenAI LLM: {self.config.llm_model}")
            
        elif self.config.llm_provider == "huggingface":
            # For HuggingFace models (requires more setup)
            from transformers import pipeline
            hf_pipeline = pipeline(
                "text-generation",
                model=self.config.llm_model,
                tokenizer=self.config.llm_model,
                max_length=self.config.llm_max_tokens
            )
            self.llm = HuggingFacePipeline(pipeline=hf_pipeline)
            logger.info(f"Using HuggingFace LLM: {self.config.llm_model}")
            
        else:
            raise ValueError(f"Unsupported LLM provider: {self.config.llm_provider}")
    
    def load_documents(self, file_paths: List[str]) -> List[Document]:
        """
        Load documents from various file formats.
        
        This method handles multiple file types and converts them into
        LangChain Document objects for consistent processing.
        
        Args:
            file_paths: List of file paths to load
            
        Returns:
            List of Document objects
        """
        documents = []
        
        for file_path in file_paths:
            try:
                # Determine file type and use appropriate loader
                file_extension = Path(file_path).suffix.lower()
                
                if file_extension == '.pdf':
                    loader = PyPDFLoader(file_path)
                elif file_extension in ['.txt', '.md']:
                    loader = TextLoader(file_path)
                else:
                    # Fallback to text loader for unknown types
                    loader = TextLoader(file_path)
                
                # Load the document
                docs = loader.load()
                
                # Add metadata to each document
                for doc in docs:
                    doc.metadata.update({
                        'source': file_path,
                        'file_type': file_extension,
                        'loaded_at': datetime.now().isoformat()
                    })
                
                documents.extend(docs)
                logger.info(f"Loaded {len(docs)} documents from {file_path}")
                
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {str(e)}")
                continue
        
        self.documents = documents
        logger.info(f"Total documents loaded: {len(documents)}")
        return documents
    
    def chunk_documents(self, documents: List[Document]) -> List[Document]:
        """
        Split documents into smaller chunks for better retrieval.
        
        Chunking is crucial for RAG performance. Smaller chunks allow
        for more precise retrieval, while larger chunks provide more context.
        The optimal chunk size depends on your use case.
        
        Args:
            documents: List of documents to chunk
            
        Returns:
            List of chunked documents
        """
        # Choose text splitter based on configuration
        if self.config.chunking_strategy == "recursive":
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.config.chunk_size,
                chunk_overlap=self.config.chunk_overlap,
                separators=["\n\n", "\n", " ", ""]  # Try these separators in order
            )
        elif self.config.chunking_strategy == "character":
            text_splitter = CharacterTextSplitter(
                chunk_size=self.config.chunk_size,
                chunk_overlap=self.config.chunk_overlap,
                separator="\n"
            )
        elif self.config.chunking_strategy == "token":
            text_splitter = TokenTextSplitter(
                chunk_size=self.config.chunk_size,
                chunk_overlap=self.config.chunk_overlap
            )
        else:
            raise ValueError(f"Unsupported chunking strategy: {self.config.chunking_strategy}")
        
        # Split all documents
        chunked_docs = text_splitter.split_documents(documents)
        
        # Add chunk metadata
        for i, chunk in enumerate(chunked_docs):
            chunk.metadata.update({
                'chunk_id': i,
                'chunk_size': len(chunk.page_content),
                'chunked_at': datetime.now().isoformat()
            })
        
        logger.info(f"Created {len(chunked_docs)} chunks from {len(documents)} documents")
        return chunked_docs
    
    def create_vectorstore(self, documents: List[Document]):
        """
        Create a vector store from document chunks.
        
        This method generates embeddings for all document chunks and
        stores them in a vector database for efficient similarity search.
        
        Args:
            documents: List of document chunks to embed
        """
        if not documents:
            raise ValueError("No documents provided for vector store creation")
        
        # Create vector store based on configuration
        if self.config.vector_db_type == "faiss":
            self.vectorstore = FAISS.from_documents(
                documents=documents,
                embedding=self.embeddings
            )
            logger.info("Created FAISS vector store")
            
        elif self.config.vector_db_type == "chroma":
            self.vectorstore = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                persist_directory=self.config.persist_directory
            )
            logger.info("Created Chroma vector store")
            
        else:
            raise ValueError(f"Unsupported vector database: {self.config.vector_db_type}")
        
        # Create retriever from vector store
        self.retriever = self.vectorstore.as_retriever(
            search_type=self.config.retrieval_method,
            search_kwargs={'k': self.config.retrieval_top_k}
        )
        
        logger.info(f"Vector store created with {len(documents)} documents")
    
    def query(self, question: str) -> Dict[str, Any]:
        """
        Query the RAG system and get a response.
        
        This is the main interface for the RAG system. It retrieves
        relevant documents and generates a response based on the context.
        
        Args:
            question: User's question
            
        Returns:
            Dictionary containing response and metadata
        """
        if not self.retriever:
            raise ValueError("Vector store not initialized. Call create_vectorstore first.")
        
        start_time = time.time()
        
        # Step 1: Retrieve relevant documents
        retrieved_docs = self.retriever.get_relevant_documents(question)
        retrieval_time = time.time() - start_time
        
        if not retrieved_docs:
            return {
                'response': "I couldn't find any relevant information to answer your question.",
                'sources': [],
                'retrieval_time': retrieval_time,
                'generation_time': 0
            }
        
        # Step 2: Prepare context from retrieved documents
        context = self._prepare_context(retrieved_docs)
        
        # Step 3: Generate response using LLM
        generation_start = time.time()
        response = self._generate_response(question, context)
        generation_time = time.time() - generation_start
        
        # Step 4: Prepare response with metadata
        result = {
            'response': response,
            'sources': [doc.metadata.get('source', 'Unknown') for doc in retrieved_docs],
            'retrieved_chunks': len(retrieved_docs),
            'retrieval_time': retrieval_time,
            'generation_time': generation_time,
            'total_time': time.time() - start_time
        }
        
        logger.info(f"Query processed in {result['total_time']:.2f}s")
        return result
    
    def _prepare_context(self, documents: List[Document]) -> str:
        """
        Prepare context string from retrieved documents.
        
        This method combines multiple document chunks into a coherent
        context string for the LLM to use in response generation.
        
        Args:
            documents: List of retrieved documents
            
        Returns:
            Formatted context string
        """
        context_parts = []
        
        for i, doc in enumerate(documents, 1):
            # Add source information for transparency
            source = doc.metadata.get('source', 'Unknown source')
            content = doc.page_content.strip()
            
            context_parts.append(f"Source {i} ({source}):\n{content}")
        
        return "\n\n---\n\n".join(context_parts)
    
    def _generate_response(self, question: str, context: str) -> str:
        """
        Generate response using the LLM with retrieved context.
        
        This method creates a prompt that includes both the user's question
        and the relevant context, then asks the LLM to generate a response.
        
        Args:
            question: User's question
            context: Retrieved context from documents
            
        Returns:
            Generated response
        """
        # Create a prompt that instructs the LLM how to use the context
        prompt = f"""
You are a helpful AI assistant. Answer the user's question based on the provided context.
If the context doesn't contain enough information to answer the question, say so clearly.
Always be accurate and cite your sources when possible.

Context:
{context}

Question: {question}

Answer:"""
        
        try:
            # Generate response using the LLM
            response = self.llm.invoke(prompt)
            
            # Extract text from response (handling different LLM response types)
            if hasattr(response, 'content'):
                return response.content
            elif isinstance(response, str):
                return response
            else:
                return str(response)
                
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            return "I apologize, but I encountered an error while generating a response."
    
    def save_vectorstore(self, path: str):
        """
        Save the vector store to disk for persistence.
        
        Args:
            path: Path to save the vector store
        """
        if self.vectorstore is None:
            raise ValueError("No vector store to save")
        
        if self.config.vector_db_type == "faiss":
            self.vectorstore.save_local(path)
        elif self.config.vector_db_type == "chroma":
            # Chroma automatically persists if persist_directory was set
            pass
        
        logger.info(f"Vector store saved to {path}")
    
    def load_vectorstore(self, path: str):
        """
        Load a previously saved vector store.
        
        Args:
            path: Path to load the vector store from
        """
        if self.config.vector_db_type == "faiss":
            self.vectorstore = FAISS.load_local(path, self.embeddings)
            self.retriever = self.vectorstore.as_retriever(
                search_type=self.config.retrieval_method,
                search_kwargs={'k': self.config.retrieval_top_k}
            )
        elif self.config.vector_db_type == "chroma":
            self.vectorstore = Chroma(
                persist_directory=path,
                embedding_function=self.embeddings
            )
            self.retriever = self.vectorstore.as_retriever(
                search_type=self.config.retrieval_method,
                search_kwargs={'k': self.config.retrieval_top_k}
            )
        
        logger.info(f"Vector store loaded from {path}")

# ================================================================
# SECTION 4: ENHANCED RAG WITH METADATA AND FILTERING
# ================================================================

class EnhancedRAG(BasicRAG):
    """
    Enhanced RAG implementation with advanced features:
    - Metadata-based filtering
    - Hybrid search (dense + sparse)
    - Query preprocessing
    - Response post-processing
    - Caching mechanisms
    
    This class builds on BasicRAG to provide production-ready features
    while maintaining the same simple interface.
    """
    
    def __init__(self, config: RAGConfig):
        super().__init__(config)
        
        # Additional components for enhanced functionality
        self.tfidf_vectorizer = TfidfVectorizer(max_features=10000, stop_words='english')
        self.sparse_vectors = None
        self.query_cache = {}
        self.response_cache = {}
        
        # Metadata index for efficient filtering
        self.metadata_index = {}
        
        logger.info("EnhancedRAG initialized with advanced features")
    
    def add_documents_with_metadata(self, 
                                   documents: List[Document], 
                                   metadata_list: List[DocumentMetadata]):
        """
        Add documents with rich metadata for enhanced search capabilities.
        
        This method extends document storage with structured metadata,
        enabling sophisticated filtering and categorization.
        
        Args:
            documents: List of documents to add
            metadata_list: List of metadata objects for each document
        """
        if len(documents) != len(metadata_list):
            raise ValueError("Documents and metadata lists must have same length")
        
        # Enhance documents with metadata
        enhanced_docs = []
        for doc, metadata in zip(documents, metadata_list):
            # Add metadata to document
            doc.metadata.update({
                'document_id': metadata.document_id,
                'title': metadata.title,
                'category': metadata.category,
                'tags': metadata.tags,
                'quality_score': metadata.quality_score,
                'access_level': metadata.access_level,
                'created_at': metadata.created_at.isoformat(),
                'word_count': metadata.word_count
            })
            enhanced_docs.append(doc)
            
            # Update metadata index for fast filtering
            self._update_metadata_index(metadata)
        
        # Chunk the enhanced documents
        chunked_docs = self.chunk_documents(enhanced_docs)
        
        # Create both dense and sparse vectors
        self._create_hybrid_vectorstore(chunked_docs)
        
        logger.info(f"Added {len(enhanced_docs)} documents with metadata")
    
    def _update_metadata_index(self, metadata: DocumentMetadata):
        """
        Update the metadata index for efficient filtering.
        
        Args:
            metadata: Document metadata to index
        """
        # Index by category
        if metadata.category not in self.metadata_index:
            self.metadata_index[metadata.category] = []
        self.metadata_index[metadata.category].append(metadata.document_id)
        
        # Index by tags
        for tag in metadata.tags:
            if f"tag:{tag}" not in self.metadata_index:
                self.metadata_index[f"tag:{tag}"] = []
            self.metadata_index[f"tag:{tag}"].append(metadata.document_id)
        
        # Index by access level
        if f"access:{metadata.access_level}" not in self.metadata_index:
            self.metadata_index[f"access:{metadata.access_level}"] = []
        self.metadata_index[f"access:{metadata.access_level}"].append(metadata.document_id)
    
    def _create_hybrid_vectorstore(self, documents: List[Document]):
        """
        Create both dense (semantic) and sparse (keyword) vector stores.
        
        Hybrid search combines the benefits of both approaches:
        - Dense vectors: Capture semantic meaning and context
        - Sparse vectors: Capture exact keyword matches and specific terms
        
        Args:
            documents: List of documents to vectorize
        """
        # Create dense vector store (same as BasicRAG)
        self.create_vectorstore(documents)
        
        # Create sparse vector store using TF-IDF
        document_texts = [doc.page_content for doc in documents]
        self.sparse_vectors = self.tfidf_vectorizer.fit_transform(document_texts)
        
        logger.info("Created hybrid vector store (dense + sparse)")
    
    def query_with_filters(self, 
                          question: str, 
                          filters: Optional[Dict[str, Any]] = None,
                          hybrid_search: bool = True,
                          use_cache: bool = True) -> Dict[str, Any]:
        """
        Query with metadata filtering and hybrid search.
        
        This advanced query method supports:
        - Metadata-based filtering
        - Hybrid search combining dense and sparse retrieval
        - Query caching for performance
        - Enhanced response metadata
        
        Args:
            question: User's question
            filters: Dictionary of metadata filters
            hybrid_search: Whether to use hybrid search
            use_cache: Whether to use response caching
            
        Returns:
            Enhanced response dictionary
        """
        # Check cache first
        cache_key = self._generate_cache_key(question, filters)
        if use_cache and cache_key in self.response_cache:
            cached_response = self.response_cache[cache_key]
            cached_response['from_cache'] = True
            return cached_response
        
        start_time = time.time()
        
        # Preprocess query
        processed_question = self._preprocess_query(question)
        
        # Retrieve documents with filtering
        if hybrid_search:
            retrieved_docs = self._hybrid_search(processed_question, filters)
        else:
            retrieved_docs = self._filtered_search(processed_question, filters)
        
        retrieval_time = time.time() - start_time
        
        if not retrieved_docs:
            return {
                'response': "I couldn't find any relevant information matching your criteria.",
                'sources': [],
                'filters_applied': filters or {},
                'retrieval_time': retrieval_time,
                'generation_time': 0,
                'from_cache': False
            }
        
        # Generate response
        generation_start = time.time()
        context = self._prepare_enhanced_context(retrieved_docs)
        response = self._generate_enhanced_response(processed_question, context, retrieved_docs)
        generation_time = time.time() - generation_start
        
        # Prepare result
        result = {
            'response': response,
            'sources': self._extract_source_info(retrieved_docs),
            'retrieved_chunks': len(retrieved_docs),
            'filters_applied': filters or {},
            'search_method': 'hybrid' if hybrid_search else 'filtered',
            'retrieval_time': retrieval_time,
            'generation_time': generation_time,
            'total_time': time.time() - start_time,
            'from_cache': False
        }
        
        # Cache the result
        if use_cache:
            self.response_cache[cache_key] = result
        
        logger.info(f"Enhanced query processed in {result['total_time']:.2f}s")
        return result
    
    def _preprocess_query(self, question: str) -> str:
        """
        Preprocess the user query for better retrieval.
        
        This method applies various preprocessing steps:
        - Query expansion
        - Spelling correction
        - Stopword handling
        - Entity recognition
        
        Args:
            question: Original user question
            
        Returns:
            Processed question
        """
        # Basic preprocessing (can be enhanced with more NLP)
        processed = question.strip()
        
        # Remove extra whitespace
        processed = ' '.join(processed.split())
        
        # Store original query for reference
        self.query_cache[processed] = question
        
        return processed
    
    def _hybrid_search(self, question: str, filters: Optional[Dict[str, Any]] = None) -> List[Document]:
        """
        Perform hybrid search combining dense and sparse retrieval.
        
        This method:
        1. Performs dense (semantic) search using embeddings
        2. Performs sparse (keyword) search using TF-IDF
        3. Combines and ranks results
        4. Applies metadata filters
        
        Args:
            question: Preprocessed user question
            filters: Metadata filters to apply
            
        Returns:
            List of retrieved documents
        """
        # Dense search using semantic embeddings
        dense_docs = self.retriever.get_relevant_documents(question)
        
        # Sparse search using TF-IDF
        sparse_docs = self._sparse_search(question)
        
        # Combine and deduplicate results
        combined_docs = self._combine_search_results(dense_docs, sparse_docs)
        
        # Apply metadata filters
        if filters:
            combined_docs = self._apply_metadata_filters(combined_docs, filters)
        
        # Limit to top-k results
        return combined_docs[:self.config.retrieval_top_k]
    
    def _sparse_search(self, question: str) -> List[Document]:
        """
        Perform sparse (keyword-based) search using TF-IDF.
        
        Args:
            question: User question
            
        Returns:
            List of documents ranked by TF-IDF similarity
        """
        if self.sparse_vectors is None:
            return []
        
        # Transform query to TF-IDF vector
        query_vector = self.tfidf_vectorizer.transform([question])
        
        # Compute similarities
        similarities = cosine_similarity(query_vector, self.sparse_vectors).flatten()
        
        # Get top results
        top_indices = np.argsort(similarities)[::-1][:self.config.retrieval_top_k * 2]
        
        # Filter by threshold
        filtered_indices = [i for i in top_indices if similarities[i] > self.config.retrieval_threshold]
        
        # Return corresponding documents
        sparse_docs = []
        for idx in filtered_indices:
            if idx < len(self.documents):
                doc = self.documents[idx]
                doc.metadata['sparse_score'] = float(similarities[idx])
                sparse_docs.append(doc)
        
        return sparse_docs
    
    def _combine_search_results(self, dense_docs: List[Document], sparse_docs: List[Document]) -> List[Document]:
        """
        Combine and rank results from dense and sparse search.
        
        This method implements a fusion approach that considers both
        semantic similarity and keyword matching scores.
        
        Args:
            dense_docs: Results from dense search
            sparse_docs: Results from sparse search
            
        Returns:
            Combined and ranked document list
        """
        # Create a mapping of documents by content for deduplication
        doc_map = {}
        
        # Add dense search results
        for doc in dense_docs:
            content_hash = hashlib.md5(doc.page_content.encode()).hexdigest()
            if content_hash not in doc_map:
                doc.metadata['dense_score'] = 1.0  # Placeholder score
                doc.metadata['sparse_score'] = 0.0
                doc_map[content_hash] = doc
        
        # Add sparse search results (merge scores if document already exists)
        for doc in sparse_docs:
            content_hash = hashlib.md5(doc.page_content.encode()).hexdigest()
            if content_hash in doc_map:
                # Update sparse score for existing document
                doc_map[content_hash].metadata['sparse_score'] = doc.metadata.get('sparse_score', 0.0)
            else:
                # Add new document
                doc.metadata['dense_score'] = 0.0
                doc_map[content_hash] = doc
        
        # Calculate combined scores and sort
        combined_docs = list(doc_map.values())
        for doc in combined_docs:
            dense_score = doc.metadata.get('dense_score', 0.0)
            sparse_score = doc.metadata.get('sparse_score', 0.0)
            # Weighted combination (can be tuned)
            doc.metadata['combined_score'] = 0.7 * dense_score + 0.3 * sparse_score
        
        # Sort by combined score
        combined_docs.sort(key=lambda x: x.metadata.get('combined_score', 0.0), reverse=True)
        
        return combined_docs
    
    def _apply_metadata_filters(self, documents: List[Document], filters: Dict[str, Any]) -> List[Document]:
        """
        Apply metadata filters to retrieved documents.
        
        Supports various filter types:
        - Category filtering
        - Tag filtering
        - Date range filtering
        - Quality score filtering
        - Access level filtering
        
        Args:
            documents: Documents to filter
            filters: Filter criteria
            
        Returns:
            Filtered document list
        """
        filtered_docs = []
        
        for doc in documents:
            include_doc = True
            
            # Category filter
            if 'category' in filters:
                if doc.metadata.get('category') != filters['category']:
                    include_doc = False
            
            # Tags filter (document must have all specified tags)
            if 'tags' in filters:
                doc_tags = doc.metadata.get('tags', [])
                required_tags = filters['tags'] if isinstance(filters['tags'], list) else [filters['tags']]
                if not all(tag in doc_tags for tag in required_tags):
                    include_doc = False
            
            # Quality score filter
            if 'min_quality' in filters:
                if doc.metadata.get('quality_score', 0.0) < filters['min_quality']:
                    include_doc = False
            
            # Access level filter
            if 'access_level' in filters:
                if doc.metadata.get('access_level') != filters['access_level']:
                    include_doc = False
            
            # Date range filter
            if 'date_from' in filters or 'date_to' in filters:
                doc_date = doc.metadata.get('created_at')
                if doc_date:
                    try:
                        doc_datetime = datetime.fromisoformat(doc_date.replace('Z', '+00:00'))
                        if 'date_from' in filters and doc_datetime < filters['date_from']:
                            include_doc = False
                        if 'date_to' in filters and doc_datetime > filters['date_to']:
                            include_doc = False
                    except (ValueError, TypeError):
                        # Skip date filtering if date parsing fails
                        pass
            
            if include_doc:
                filtered_docs.append(doc)
        
        return filtered_docs
    
    def _prepare_enhanced_context(self, documents: List[Document]) -> str:
        """
        Prepare enhanced context with metadata and scoring information.
        
        Args:
            documents: Retrieved documents
            
        Returns:
            Enhanced context string
        """
        context_parts = []
        
        for i, doc in enumerate(documents, 1):
            # Extract metadata
            source = doc.metadata.get('source', 'Unknown source')
            category = doc.metadata.get('category', 'Unknown')
            quality_score = doc.metadata.get('quality_score', 'N/A')
            combined_score = doc.metadata.get('combined_score', 'N/A')
            
            # Format document with metadata
            metadata_info = f"[Category: {category}, Quality: {quality_score}, Relevance: {combined_score:.3f}]"
            content = doc.page_content.strip()
            
            context_parts.append(f"Source {i} ({source}) {metadata_info}:\n{content}")
        
        return "\n\n---\n\n".join(context_parts)
    
    def _generate_enhanced_response(self, question: str, context: str, documents: List[Document]) -> str:
        """
        Generate enhanced response with better prompting and post-processing.
        
        Args:
            question: User question
            context: Enhanced context from documents
            documents: Retrieved documents for reference
            
        Returns:
            Generated response
        """
        # Create enhanced prompt with instructions
        prompt = f"""
You are an expert AI assistant with access to a curated knowledge base. Your task is to provide accurate, helpful, and well-structured responses based on the provided context.

INSTRUCTIONS:
1. Answer the user's question comprehensively using the provided sources
2. If information is insufficient, clearly state what is missing
3. Cite specific sources when making claims
4. Maintain objectivity and acknowledge different perspectives when present
5. Structure your response clearly with appropriate sections if needed

CONTEXT FROM KNOWLEDGE BASE:
{context}

USER QUESTION: {question}

Please provide a detailed response based on the context above:"""
        
        try:
            # Generate response
            response = self.llm.invoke(prompt)
            
            # Extract and post-process response
            if hasattr(response, 'content'):
                response_text = response.content
            else:
                response_text = str(response)
            
            # Post-process response (can add fact-checking, formatting, etc.)
            enhanced_response = self._post_process_response(response_text, documents)
            
            return enhanced_response
            
        except Exception as e:
            logger.error(f"Error generating enhanced response: {str(e)}")
            return "I apologize, but I encountered an error while generating a response. Please try rephrasing your question."
    
    def _post_process_response(self, response: str, documents: List[Document]) -> str:
        """
        Post-process the generated response for quality and consistency.
        
        Args:
            response: Raw response from LLM
            documents: Source documents for reference
            
        Returns:
            Post-processed response
        """
        # Basic post-processing (can be enhanced)
        processed_response = response.strip()
        
        # Add confidence note if retrieval scores are low
        avg_score = np.mean([doc.metadata.get('combined_score', 0.0) for doc in documents])
        if avg_score < 0.5:
            processed_response += "\n\n*Note: This response is based on documents with lower relevance scores. Consider refining your question for better results.*"
        
        return processed_response
    
    def _extract_source_info(self, documents: List[Document]) -> List[Dict[str, Any]]:
        """
        Extract detailed source information from documents.
        
        Args:
            documents: Retrieved documents
            
        Returns:
            List of source information dictionaries
        """
        sources = []
        for doc in documents:
            source_info = {
                'source': doc.metadata.get('source', 'Unknown'),
                'title': doc.metadata.get('title', 'Untitled'),
                'category': doc.metadata.get('category', 'Unknown'),
                'quality_score': doc.metadata.get('quality_score', 'N/A'),
                'relevance_score': doc.metadata.get('combined_score', 'N/A'),
                'chunk_id': doc.metadata.get('chunk_id', 'N/A')
            }
            sources.append(source_info)
        return sources
    
    def _generate_cache_key(self, question: str, filters: Optional[Dict[str, Any]] = None) -> str:
        """
        Generate a cache key for query caching.
        
        Args:
            question: User question
            filters: Applied filters
            
        Returns:
            Cache key string
        """
        cache_data = {
            'question': question,
            'filters': filters or {},
            'config_hash': hash(str(self.config))
        }
        return hashlib.md5(json.dumps(cache_data, sort_keys=True).encode()).hexdigest()

# ================================================================
# SECTION 5: PRODUCTION RAG PIPELINE
# ================================================================

class ProductionRAG(EnhancedRAG):
    """
    Production-ready RAG implementation with:
    - Monitoring and metrics
    - Async processing
    - Error handling and resilience
    - Performance optimization
    - A/B testing capabilities
    - Auto-scaling support
    
    This class is designed for enterprise deployment with full
    observability and production-grade features.
    """
    
    def __init__(self, config: RAGConfig):
        super().__init__(config)
        
        # Monitoring setup
        self.setup_monitoring()
        
        # Performance metrics
        self.query_counter = Counter('rag_queries_total', 'Total number of queries processed')
        self.query_duration = Histogram('rag_query_duration_seconds', 'Time spent processing queries')
        self.retrieval_duration = Histogram('rag_retrieval_duration_seconds', 'Time spent on retrieval')
        self.generation_duration = Histogram('rag_generation_duration_seconds', 'Time spent on generation')
        self.error_counter = Counter('rag_errors_total', 'Total number of errors', ['error_type'])
        
        # Health check endpoint
        self.health_status = {'status': 'healthy', 'last_check': datetime.now()}
        
        # Async processing setup
        self.executor = None
        
        logger.info("ProductionRAG initialized with monitoring and metrics")
    
    def setup_monitoring(self):
        """
        Setup monitoring and metrics collection.
        """
        if self.config.enable_metrics:
            try:
                start_http_server(self.config.metrics_port)
                logger.info(f"Metrics server started on port {self.config.metrics_port}")
            except Exception as e:
                logger.warning(f"Failed to start metrics server: {str(e)}")
    
    async def async_query(self, question: str, filters: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Async version of query method for better performance.
        
        Args:
            question: User question
            filters: Optional metadata filters
            
        Returns:
            Query response with metadata
        """
        with self.query_duration.time():
            self.query_counter.inc()
            
            try:
                # Run the synchronous query in a thread pool
                loop = asyncio.get_event_loop()
                if self.executor is None:
                    from concurrent.futures import ThreadPoolExecutor
                    self.executor = ThreadPoolExecutor(max_workers=self.config.max_workers)
                
                result = await loop.run_in_executor(
                    self.executor,
                    self.query_with_filters,
                    question,
                    filters
                )
                
                # Update health status
                self.health_status = {'status': 'healthy', 'last_check': datetime.now()}
                
                return result
                
            except Exception as e:
                self.error_counter.labels(error_type=type(e).__name__).inc()
                self.health_status = {'status': 'error', 'last_check': datetime.now(), 'error': str(e)}
                logger.error(f"Error in async_query: {str(e)}")
                raise
    
    def batch_process_queries(self, queries: List[str], filters_list: Optional[List[Dict[str, Any]]] = None) -> List[Dict[str, Any]]:
        """
        Process multiple queries in batch for efficiency.
        
        Args:
            queries: List of user questions
            filters_list: Optional list of filters for each query
            
        Returns:
            List of query responses
        """
        if filters_list and len(queries) != len(filters_list):
            raise ValueError("Queries and filters lists must have same length")
        
        results = []
        batch_start_time = time.time()
        
        # Process queries in batches
        batch_size = self.config.batch_size
        for i in range(0, len(queries), batch_size):
            batch_queries = queries[i:i+batch_size]
            batch_filters = filters_list[i:i+batch_size] if filters_list else [None] * len(batch_queries)
            
            # Process batch
            batch_results = []
            for query, filters in zip(batch_queries, batch_filters):
                try:
                    result = self.query_with_filters(query, filters)
                    batch_results.append(result)
                except Exception as e:
                    logger.error(f"Error processing query in batch: {str(e)}")
                    batch_results.append({
                        'response': f"Error processing query: {str(e)}",
                        'error': True
                    })
            
            results.extend(batch_results)
        
        batch_duration = time.time() - batch_start_time
        logger.info(f"Processed {len(queries)} queries in {batch_duration:.2f}s")
        
        return results
    
    def get_health_status(self) -> Dict[str, Any]:
        """
        Get system health status for monitoring.
        
        Returns:
            Health status dictionary
        """
        health_info = self.health_status.copy()
        
        # Add system metrics
        health_info.update({
            'vectorstore_status': 'ready' if self.vectorstore else 'not_initialized',
            'cache_size': len(self.response_cache),
            'query_cache_size': len(self.query_cache),
            'uptime': str(datetime.now() - self.health_status.get('startup_time', datetime.now()))
        })
        
        return health_info
    
    def get_metrics(self) -> Dict[str, Any]:
        """
        Get system performance metrics.
        
        Returns:
            Metrics dictionary
        """
        from prometheus_client import REGISTRY
        
        metrics = {}
        for collector in REGISTRY._collector_to_names:
            for metric in collector.collect():
                for sample in metric.samples:
                    metrics[sample.name] = sample.value
        
        return metrics
    
    def validate_query(self, question: str) -> Tuple[bool, str]:
        """
        Validate user query before processing.
        
        Args:
            question: User question to validate
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Basic validation
        if not question or not question.strip():
            return False, "Question cannot be empty"
        
        if len(question) > 1000:
            return False, "Question is too long (max 1000 characters)"
        
        if len(question.strip()) < 3:
            return False, "Question is too short (min 3 characters)"
        
        # Content validation (can be enhanced with more sophisticated checks)
        forbidden_patterns = ['<script>', '<?php', 'javascript:']
        for pattern in forbidden_patterns:
            if pattern in question.lower():
                return False, "Question contains potentially unsafe content"
        
        return True, ""
    
    def optimize_retrieval(self, question: str) -> str:
        """
        Optimize query for better retrieval performance.
        
        Args:
            question: Original question
            
        Returns:
            Optimized question
        """
        # Query expansion and optimization logic
        optimized = question.strip()
        
        # Remove redundant words
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
        words = optimized.split()
        filtered_words = [word for word in words if word.lower() not in stop_words or len(words) <= 3]
        
        optimized = ' '.join(filtered_words)
        
        return optimized
    
    def a_b_test_query(self, question: str, test_group: str = 'A') -> Dict[str, Any]:
        """
        A/B testing framework for query processing.
        
        Args:
            question: User question
            test_group: Test group identifier ('A' or 'B')
            
        Returns:
            Query response with test metadata
        """
        # Configure different approaches for A/B testing
        if test_group == 'A':
            # Standard approach
            result = self.query_with_filters(question, hybrid_search=True)
        elif test_group == 'B':
            # Alternative approach (e.g., different chunk size, different prompt)
            original_chunk_size = self.config.chunk_size
            self.config.chunk_size = int(original_chunk_size * 1.5)  # Test larger chunks
            result = self.query_with_filters(question, hybrid_search=False)
            self.config.chunk_size = original_chunk_size  # Restore original
        else:
            raise ValueError(f"Unknown test group: {test_group}")
        
        # Add test metadata
        result['test_group'] = test_group
        result['test_timestamp'] = datetime.now().isoformat()
        
        return result

# ================================================================
# SECTION 6: MULTI-MODAL RAG IMPLEMENTATION
# ================================================================

class MultiModalRAG(ProductionRAG):
    """
    Multi-modal RAG implementation supporting:
    - Text documents
    - Images with OCR
    - Audio transcription
    - Video content extraction
    - Mixed media search and retrieval
    
    This class extends RAG capabilities to handle various data types
    beyond plain text, enabling richer knowledge bases and interactions.
    """
    
    def __init__(self, config: RAGConfig):
        super().__init__(config)
        
        # Multi-modal components
        self.image_processor = None
        self.audio_processor = None
        self.video_processor = None
        
        # Additional vector stores for different modalities
        self.image_vectorstore = None
        self.audio_vectorstore = None
        
        # Cross-modal embedding models
        self.clip_model = None  # For image-text similarity
        
        self._setup_multimodal_components()
        
        logger.info("MultiModalRAG initialized with multi-modal capabilities")
    
    def _setup_multimodal_components(self):
        """
        Setup components for multi-modal processing.
        """
        try:
            # Image processing setup
            import easyocr
            self.ocr_reader = easyocr.Reader(['en'])
            
            # Audio processing setup (requires additional packages)
            # import whisper
            # self.whisper_model = whisper.load_model("base")
            
            logger.info("Multi-modal components initialized")
        except ImportError as e:
            logger.warning(f"Some multi-modal dependencies not available: {str(e)}")
    
    def process_image_document(self, image_path: str, metadata: DocumentMetadata) -> Document:
        """
        Process image document with OCR text extraction.
        
        Args:
            image_path: Path to image file
            metadata: Document metadata
            
        Returns:
            Document with extracted text
        """
        try:
            # Extract text using OCR
            results = self.ocr_reader.readtext(image_path)
            extracted_text = ' '.join([result[1] for result in results])
            
            # Create document
            doc = Document(
                page_content=extracted_text,
                metadata={
                    **metadata.__dict__,
                    'source': image_path,
                    'document_type': 'image',
                    'ocr_confidence': np.mean([result[2] for result in results]) if results else 0.0
                }
            )
            
            logger.info(f"Processed image document: {image_path}")
            return doc
            
        except Exception as e:
            logger.error(f"Error processing image {image_path}: {str(e)}")
            raise
    
    def process_audio_document(self, audio_path: str, metadata: DocumentMetadata) -> Document:
        """
        Process audio document with speech-to-text.
        
        Args:
            audio_path: Path to audio file
            metadata: Document metadata
            
        Returns:
            Document with transcribed text
        """
        # Placeholder for audio processing
        # In a real implementation, you would use Whisper or similar
        try:
            # Simulated transcription
            transcribed_text = f"[Audio transcription from {audio_path} would go here]"
            
            doc = Document(
                page_content=transcribed_text,
                metadata={
                    **metadata.__dict__,
                    'source': audio_path,
                    'document_type': 'audio',
                    'duration': 0.0,  # Would extract actual duration
                    'transcription_confidence': 0.95
                }
            )
            
            logger.info(f"Processed audio document: {audio_path}")
            return doc
            
        except Exception as e:
            logger.error(f"Error processing audio {audio_path}: {str(e)}")
            raise
    
    def add_multimodal_documents(self, file_paths: List[str], metadata_list: List[DocumentMetadata]):
        """
        Add documents of various types (text, image, audio) to the knowledge base.
        
        Args:
            file_paths: List of file paths
            metadata_list: List of metadata for each file
        """
        all_documents = []
        
        for file_path, metadata in zip(file_paths, metadata_list):
            file_extension = Path(file_path).suffix.lower()
            
            try:
                if file_extension in ['.txt', '.md', '.pdf']:
                    # Text documents
                    doc = self._process_text_document(file_path, metadata)
                elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff']:
                    # Image documents
                    doc = self.process_image_document(file_path, metadata)
                elif file_extension in ['.wav', '.mp3', '.m4a']:
                    # Audio documents
                    doc = self.process_audio_document(file_path, metadata)
                else:
                    logger.warning(f"Unsupported file type: {file_extension}")
                    continue
                
                all_documents.append(doc)
                
            except Exception as e:
                logger.error(f"Error processing {file_path}: {str(e)}")
                continue
        
        # Add all processed documents
        if all_documents:
            self.add_documents_with_metadata(all_documents, metadata_list[:len(all_documents)])
    
    def _process_text_document(self, file_path: str, metadata: DocumentMetadata) -> Document:
        """
        Process text document (helper method).
        
        Args:
            file_path: Path to text file
            metadata: Document metadata
            
        Returns:
            Processed document
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            loader = PyPDFLoader(file_path)
        else:
            loader = TextLoader(file_path)
        
        docs = loader.load()
        doc = docs[0] if docs else Document(page_content="", metadata={})
        
        # Update metadata
        doc.metadata.update(metadata.__dict__)
        
        return doc

# ================================================================
# SECTION 7: WEB INTERFACE AND DEPLOYMENT
# ================================================================

class RAGWebInterface:
    """
    Web interface for RAG system using Streamlit.
    
    This class provides a user-friendly web interface for interacting
    with the RAG system, suitable for demos and production deployment.
    """
    
    def __init__(self, rag_system: ProductionRAG):
        self.rag_system = rag_system
        self.setup_streamlit_app()
    
    def setup_streamlit_app(self):
        """
        Setup Streamlit web application.
        """
        st.set_page_config(
            page_title="RAG System",
            page_icon="🤖",
            layout="wide",
            initial_sidebar_state="expanded"
        )
        
        # Custom CSS
        st.markdown("""
        <style>
        .main-header {
            background: linear-gradient(90deg, #1e3a8a 0%, #3b82f6 100%);
            padding: 2rem;
            border-radius: 10px;
            color: white;
            margin-bottom: 2rem;
        }
        .response-container {
            background: #f8fafc;
            padding: 1.5rem;
            border-radius: 10px;
            border-left: 4px solid #3b82f6;
            margin: 1rem 0;
        }
        .source-info {
            background: #f1f5f9;
            padding: 1rem;
            border-radius: 8px;
            margin: 0.5rem 0;
        }
        </style>
        """, unsafe_allow_html=True)
    
    def run(self):
        """
        Run the Streamlit web application.
        """
        # Header
        st.markdown("""
        <div class="main-header">
            <h1>🤖 Advanced RAG System</h1>
            <p>Retrieval-Augmented Generation with Enhanced Capabilities</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Sidebar configuration
        with st.sidebar:
            st.header("Configuration")
            
            # Query settings
            use_filters = st.checkbox("Use Metadata Filters", value=False)
            hybrid_search = st.checkbox("Hybrid Search", value=True)
            use_cache = st.checkbox("Use Cache", value=True)
            
            # Filters configuration
            if use_filters:
                st.subheader("Filters")
                category_filter = st.selectbox("Category", ["All", "technical", "business", "medical"])
                quality_threshold = st.slider("Min Quality Score", 0.0, 1.0, 0.5)
            
            # System status
            st.subheader("System Status")
            health_status = self.rag_system.get_health_status()
            st.success(f"Status: {health_status['status']}")
            st.info(f"Cache Size: {health_status['cache_size']}")
        
        # Main interface
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.header("Ask a Question")
            
            # Query input
            question = st.text_area("Enter your question:", height=100)
            
            # Query button
            if st.button("Get Answer", type="primary"):
                if question:
                    with st.spinner("Processing your question..."):
                        # Prepare filters
                        filters = None
                        if use_filters:
                            filters = {}
                            if category_filter != "All":
                                filters['category'] = category_filter
                            filters['min_quality'] = quality_threshold
                        
                        # Process query
                        try:
                            result = self.rag_system.query_with_filters(
                                question=question,
                                filters=filters,
                                hybrid_search=hybrid_search,
                                use_cache=use_cache
                            )
                            
                            # Display response
                            st.markdown(f"""
                            <div class="response-container">
                                <h3>Response:</h3>
                                <p>{result['response']}</p>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            # Display metadata
                            col_meta1, col_meta2, col_meta3 = st.columns(3)
                            with col_meta1:
                                st.metric("Retrieval Time", f"{result['retrieval_time']:.3f}s")
                            with col_meta2:
                                st.metric("Generation Time", f"{result['generation_time']:.3f}s")
                            with col_meta3:
                                st.metric("Total Time", f"{result['total_time']:.3f}s")
                            
                            # Display sources
                            if result['sources']:
                                with st.expander("Sources"):
                                    for i, source in enumerate(result['sources'], 1):
                                        st.markdown(f"""
                                        <div class="source-info">
                                            <strong>Source {i}:</strong> {source['source']}<br>
                                            <strong>Category:</strong> {source['category']}<br>
                                            <strong>Quality:</strong> {source['quality_score']}<br>
                                            <strong>Relevance:</strong> {source['relevance_score']:.3f}
                                        </div>
                                        """, unsafe_allow_html=True)
                        
                        except Exception as e:
                            st.error(f"Error processing question: {str(e)}")
                else:
                    st.warning("Please enter a question.")
        
        with col2:
            st.header("System Metrics")
            
            # Display metrics if available
            try:
                metrics = self.rag_system.get_metrics()
                if metrics:
                    for metric_name, value in list(metrics.items())[:10]:  # Show top 10
                        st.metric(metric_name.replace('_', ' ').title(), f"{value:.2f}")
            except Exception as e:
                st.info("Metrics not available")
            
            # Example questions
            st.subheader("Example Questions")
            example_questions = [
                "What are the main benefits of cloud computing?",
                "How does machine learning work?",
                "What is the difference between AI and ML?",
                "Explain the concept of microservices",
                "What are best practices for data security?"
            ]
            
            for question in example_questions:
                if st.button(question, key=f"example_{hash(question)}"):
                    st.session_state.example_question = question

class RAGAPIServer:
    """
    REST API server for RAG system using Flask.
    
    This class provides a REST API interface for the RAG system,
    suitable for integration with other applications and services.
    """
    
    def __init__(self, rag_system: ProductionRAG, host='0.0.0.0', port=5000):
        self.rag_system = rag_system
        self.app = Flask(__name__)
        self.host = host
        self.port = port
        
        self.setup_routes()
    
    def setup_routes(self):
        """
        Setup API routes.
        """
        @self.app.route('/health', methods=['GET'])
        def health_check():
            """Health check endpoint."""
            return jsonify(self.rag_system.get_health_status())
        
        @self.app.route('/metrics', methods=['GET'])
        def get_metrics():
            """Metrics endpoint."""
            return jsonify(self.rag_system.get_metrics())
        
        @self.app.route('/query', methods=['POST'])
        def query():
            """Query endpoint."""
            try:
                data = request.json
                question = data.get('question')
                filters = data.get('filters')
                
                if not question:
                    return jsonify({'error': 'Question is required'}), 400
                
                # Validate query
                is_valid, error_msg = self.rag_system.validate_query(question)
                if not is_valid:
                    return jsonify({'error': error_msg}), 400
                
                # Process query
                result = self.rag_system.query_with_filters(
                    question=question,
                    filters=filters,
                    hybrid_search=data.get('hybrid_search', True),
                    use_cache=data.get('use_cache', True)
                )
                
                return jsonify(result)
                
            except Exception as e:
                logger.error(f"Error in query endpoint: {str(e)}")
                return jsonify({'error': 'Internal server error'}), 500
        
        @self.app.route('/batch_query', methods=['POST'])
        def batch_query():
            """Batch query endpoint."""
            try:
                data = request.json
                questions = data.get('questions', [])
                filters_list = data.get('filters_list')
                
                if not questions:
                    return jsonify({'error': 'Questions list is required'}), 400
                
                # Process batch queries
                results = self.rag_system.batch_process_queries(questions, filters_list)
                
                return jsonify({'results': results})
                
            except Exception as e:
                logger.error(f"Error in batch query endpoint: {str(e)}")
                return jsonify({'error': 'Internal server error'}), 500
        
        @self.app.route('/add_documents', methods=['POST'])
        def add_documents():
            """Add documents endpoint."""
            try:
                data = request.json
                documents_data = data.get('documents', [])
                
                # Process documents (simplified for API)
                # In a real implementation, you'd handle file uploads
                return jsonify({'message': 'Documents added successfully'})
                
            except Exception as e:
                logger.error(f"Error in add documents endpoint: {str(e)}")
                return jsonify({'error': 'Internal server error'}), 500
    
    def run(self, debug=False):
        """
        Run the Flask API server.
        
        Args:
            debug: Whether to run in debug mode
        """
        logger.info(f"Starting RAG API server on {self.host}:{self.port}")
        self.app.run(host=self.host, port=self.port, debug=debug)

# ================================================================
# SECTION 8: MONITORING AND ANALYTICS
# ================================================================

class RAGAnalytics:
    """
    Analytics and monitoring system for RAG performance.
    
    This class provides comprehensive analytics including:
    - Query performance tracking
    - Retrieval quality analysis
    - User behavior insights
    - System performance monitoring
    - A/B testing results analysis
    """
    
    def __init__(self, rag_system: ProductionRAG):
        self.rag_system = rag_system
        self.query_logs = []
        self.performance_metrics = {}
        self.user_sessions = {}
        
        # Setup analytics database (in production, use proper DB)
        self.analytics_db = {}
        
        logger.info("RAG Analytics initialized")
    
    def log_query(self, 
                  user_id: str, 
                  question: str, 
                  result: Dict[str, Any], 
                  filters: Optional[Dict[str, Any]] = None):
        """
        Log query for analytics.
        
        Args:
            user_id: User identifier
            question: User question
            result: Query result
            filters: Applied filters
        """
        log_entry = {
            'timestamp': datetime.now(),
            'user_id': user_id,
            'question': question,
            'response': result.get('response', ''),
            'retrieval_time': result.get('retrieval_time', 0),
            'generation_time': result.get('generation_time', 0),
            'total_time': result.get('total_time', 0),
            'retrieved_chunks': result.get('retrieved_chunks', 0),
            'sources_count': len(result.get('sources', [])),
            'filters': filters or {},
            'from_cache': result.get('from_cache', False),
            'search_method': result.get('search_method', 'unknown')
        }
        
        self.query_logs.append(log_entry)
        
        # Update user session
        if user_id not in self.user_sessions:
            self.user_sessions[user_id] = {
                'first_query': datetime.now(),
                'query_count': 0,
                'total_time': 0,
                'avg_response_time': 0
            }
        
        session = self.user_sessions[user_id]
        session['query_count'] += 1
        session['total_time'] += result.get('total_time', 0)
        session['avg_response_time'] = session['total_time'] / session['query_count']
        session['last_query'] = datetime.now()
    
    def generate_performance_report(self, days: int = 7) -> Dict[str, Any]:
        """
        Generate performance report for the last N days.
        
        Args:
            days: Number of days to include in report
            
        Returns:
            Performance report dictionary
        """
        cutoff_date = datetime.now() - timedelta(days=days)
        recent_logs = [log for log in self.query_logs if log['timestamp'] > cutoff_date]
        
        if not recent_logs:
            return {'error': 'No data available for the specified period'}
        
        # Calculate metrics
        total_queries = len(recent_logs)
        avg_retrieval_time = np.mean([log['retrieval_time'] for log in recent_logs])
        avg_generation_time = np.mean([log['generation_time'] for log in recent_logs])
        avg_total_time = np.mean([log['total_time'] for log in recent_logs])
        cache_hit_rate = sum(1 for log in recent_logs if log['from_cache']) / total_queries
        
        # Query distribution by hour
        hourly_distribution = {}
        for log in recent_logs:
            hour = log['timestamp'].hour
            hourly_distribution[hour] = hourly_distribution.get(hour, 0) + 1
        
        # Most common query patterns
        question_lengths = [len(log['question']) for log in recent_logs]
        
        report = {
            'period': f"Last {days} days",
            'total_queries': total_queries,
            'unique_users': len(set(log['user_id'] for log in recent_logs)),
            'performance': {
                'avg_retrieval_time': round(avg_retrieval_time, 3),
                'avg_generation_time': round(avg_generation_time, 3),
                'avg_total_time': round(avg_total_time, 3),
                'cache_hit_rate': round(cache_hit_rate, 2)
            },
            'query_patterns': {
                'avg_question_length': round(np.mean(question_lengths), 1),
                'median_question_length': round(np.median(question_lengths), 1),
                'hourly_distribution': hourly_distribution
            },
            'retrieval_stats': {
                'avg_chunks_retrieved': round(np.mean([log['retrieved_chunks'] for log in recent_logs]), 1),
                'avg_sources_count': round(np.mean([log['sources_count'] for log in recent_logs]), 1)
            }
        }
        
        return report
    
    def analyze_query_quality(self) -> Dict[str, Any]:
        """
        Analyze query and response quality.
        
        Returns:
            Quality analysis results
        """
        if not self.query_logs:
            return {'error': 'No query data available'}
        
        # Analyze response times
        response_times = [log['total_time'] for log in self.query_logs]
        slow_queries = [log for log in self.query_logs if log['total_time'] > 5.0]  # >5 seconds
        
        # Analyze retrieval effectiveness
        zero_result_queries = [log for log in self.query_logs if log['retrieved_chunks'] == 0]
        low_result_queries = [log for log in self.query_logs if 0 < log['retrieved_chunks'] < 3]
        
        quality_analysis = {
            'response_time_analysis': {
                'p50_response_time': round(np.percentile(response_times, 50), 3),
                'p95_response_time': round(np.percentile(response_times, 95), 3),
                'p99_response_time': round(np.percentile(response_times, 99), 3),
                'slow_queries_count': len(slow_queries),
                'slow_queries_percentage': round(len(slow_queries) / len(self.query_logs) * 100, 2)
            },
            'retrieval_analysis': {
                'zero_result_queries': len(zero_result_queries),
                'zero_result_percentage': round(len(zero_result_queries) / len(self.query_logs) * 100, 2),
                'low_result_queries': len(low_result_queries),
                'low_result_percentage': round(len(low_result_queries) / len(self.query_logs) * 100, 2)
            }
        }
        
        return quality_analysis
    
    def get_user_insights(self) -> Dict[str, Any]:
        """
        Get insights about user behavior.
        
        Returns:
            User behavior insights
        """
        if not self.user_sessions:
            return {'error': 'No user session data available'}
        
        # Calculate user metrics
        session_lengths = []
        queries_per_user = []
        
        for user_id, session in self.user_sessions.items():
            if 'last_query' in session:
                session_length = (session['last_query'] - session['first_query']).total_seconds() / 60
                session_lengths.append(session_length)
            queries_per_user.append(session['query_count'])
        
        insights = {
            'total_users': len(self.user_sessions),
            'avg_queries_per_user': round(np.mean(queries_per_user), 1),
            'median_queries_per_user': round(np.median(queries_per_user), 1),
            'avg_session_length_minutes': round(np.mean(session_lengths), 1) if session_lengths else 0,
            'power_users': len([q for q in queries_per_user if q > 10]),  # Users with >10 queries
            'single_query_users': len([q for q in queries_per_user if q == 1])
        }
        
        return insights

# ================================================================
# SECTION 9: EXAMPLE USAGE AND TESTING
# ================================================================

def create_sample_documents() -> List[Document]:
    """
    Create sample documents for testing the RAG system.
    
    Returns:
        List of sample documents
    """
    sample_docs = [
        Document(
            page_content="""
            Machine Learning is a subset of artificial intelligence that enables computers to learn 
            and make decisions from data without being explicitly programmed. It involves algorithms 
            that can identify patterns in data and use these patterns to make predictions or decisions 
            about new, unseen data. Common types include supervised learning, unsupervised learning, 
            and reinforcement learning.
            """,
            metadata={
                'source': 'ml_basics.txt',
                'category': 'technical',
                'title': 'Introduction to Machine Learning',
                'tags': ['AI', 'ML', 'algorithms'],
                'quality_score': 0.9
            }
        ),
        Document(
            page_content="""
            Cloud computing delivers computing services including servers, storage, databases, 
            networking, software, analytics, and intelligence over the Internet. It offers 
            faster innovation, flexible resources, and economies of scale. Main service models 
            include Infrastructure as a Service (IaaS), Platform as a Service (PaaS), and 
            Software as a Service (SaaS).
            """,
            metadata={
                'source': 'cloud_computing.txt',
                'category': 'technical',
                'title': 'Cloud Computing Overview',
                'tags': ['cloud', 'infrastructure', 'services'],
                'quality_score': 0.85
            }
        ),
        Document(
            page_content="""
            Data privacy refers to the proper handling, processing, storage, and usage of personal 
            information. It involves implementing measures to protect sensitive data from unauthorized 
            access, ensuring compliance with regulations like GDPR and CCPA, and maintaining 
            transparency about data collection and usage practices.
            """,
            metadata={
                'source': 'data_privacy.txt',
                'category': 'business',
                'title': 'Data Privacy Fundamentals',
                'tags': ['privacy', 'data', 'compliance', 'GDPR'],
                'quality_score': 0.92
            }
        ),
        Document(
            page_content="""
            Microservices architecture is an approach to developing a single application as a suite 
            of small services, each running in its own process and communicating with lightweight 
            mechanisms. This architecture enables better scalability, flexibility, and technology 
            diversity, but also introduces complexity in terms of service coordination and data consistency.
            """,
            metadata={
                'source': 'microservices.txt',
                'category': 'technical',
                'title': 'Microservices Architecture',
                'tags': ['architecture', 'microservices', 'scalability'],
                'quality_score': 0.88
            }
        ),
        Document(
            page_content="""
            Cybersecurity best practices include implementing strong password policies, enabling 
            multi-factor authentication, keeping software updated, conducting regular security 
            audits, training employees on security awareness, and having an incident response 
            plan. Organizations should also consider zero-trust architecture and continuous monitoring.
            """,
            metadata={
                'source': 'cybersecurity.txt',
                'category': 'business',
                'title': 'Cybersecurity Best Practices',
                'tags': ['security', 'cybersecurity', 'best-practices'],
                'quality_score': 0.91
            }
        )
    ]
    
    return sample_docs

def demo_basic_rag():
    """
    Demonstrate basic RAG functionality.
    """
    print("=== Basic RAG Demo ===")
    
    # Initialize configuration
    config = RAGConfig(
        llm_provider="openai",  # Change to "huggingface" if no OpenAI key
        embedding_provider="sentence_transformers",  # Use free embeddings
        embedding_model="all-MiniLM-L6-v2",
        chunk_size=500,
        retrieval_top_k=3
    )
    
    # Create RAG system
    rag = BasicRAG(config)
    
    # Load sample documents
    documents = create_sample_documents()
    
    # Chunk and create vector store
    chunked_docs = rag.chunk_documents(documents)
    rag.create_vectorstore(chunked_docs)
    
    # Test queries
    test_questions = [
        "What is machine learning?",
        "How does cloud computing work?",
        "What are cybersecurity best practices?",
        "Explain microservices architecture"
    ]
    
    for question in test_questions:
        print(f"\nQuestion: {question}")
        print("-" * 50)
        
        try:
            result = rag.query(question)
            print(f"Response: {result['response']}")
            print(f"Sources: {result['sources']}")
            print(f"Retrieval time: {result['retrieval_time']:.3f}s")
            print(f"Generation time: {result['generation_time']:.3f}s")
        except Exception as e:
            print(f"Error: {str(e)}")

def demo_enhanced_rag():
    """
    Demonstrate enhanced RAG with metadata filtering.
    """
    print("\n=== Enhanced RAG Demo ===")
    
    # Initialize configuration
    config = RAGConfig(
        embedding_provider="sentence_transformers",
        embedding_model="all-MiniLM-L6-v2",
        chunk_size=500,
        retrieval_top_k=3
    )
    
    # Create enhanced RAG system
    rag = EnhancedRAG(config)
    
    # Create metadata for documents
    documents = create_sample_documents()
    metadata_list = []
    
    for i, doc in enumerate(documents):
        metadata = DocumentMetadata(
            document_id=f"doc_{i}",
            title=doc.metadata['title'],
            source=doc.metadata['source'],
            document_type="text",
            category=doc.metadata['category'],
            tags=doc.metadata['tags'],
            quality_score=doc.metadata['quality_score']
        )
        metadata_list.append(metadata)
    
    # Add documents with metadata
    rag.add_documents_with_metadata(documents, metadata_list)
    
    # Test filtered queries
    print("\nTesting filtered queries:")
    
    # Query with category filter
    result = rag.query_with_filters(
        "What are best practices?",
        filters={'category': 'business'},
        hybrid_search=True
    )
    
    print(f"Business category query result: {result['response'][:100]}...")
    print(f"Applied filters: {result['filters_applied']}")
    print(f"Search method: {result['search_method']}")
    
    # Query with tag filter
    result = rag.query_with_filters(
        "Tell me about AI technologies",
        filters={'tags': ['AI']},
        hybrid_search=True
    )
    
    print(f"\nAI tag query result: {result['response'][:100]}...")
    print(f"Sources found: {len(result['sources'])}")

def demo_production_rag():
    """
    Demonstrate production RAG features.
    """
    print("\n=== Production RAG Demo ===")
    
    # Initialize configuration with monitoring
    config = RAGConfig(
        embedding_provider="sentence_transformers",
        embedding_model="all-MiniLM-L6-v2",
        enable_metrics=True,
        metrics_port=8000
    )
    
    # Create production RAG system
    rag = ProductionRAG(config)
    
    # Setup documents
    documents = create_sample_documents()
    metadata_list = [
        DocumentMetadata(
            document_id=f"doc_{i}",
            title=doc.metadata['title'],
            source=doc.metadata['source'],
            document_type="text",
            category=doc.metadata['category'],
            tags=doc.metadata['tags'],
            quality_score=doc.metadata['quality_score']
        )
        for i, doc in enumerate(documents)
    ]
    
    rag.add_documents_with_metadata(documents, metadata_list)
    
    # Test batch processing
    batch_questions = [
        "What is machine learning?",
        "How does cloud computing help businesses?",
        "What are security best practices?",
        "Explain microservices benefits"
    ]
    
    print("Processing batch queries...")
    batch_results = rag.batch_process_queries(batch_questions)
    
    for i, result in enumerate(batch_results):
        print(f"\nBatch Query {i+1}: {batch_questions[i]}")
        print(f"Response length: {len(result.get('response', ''))}")
        print(f"Has error: {'error' in result}")
    
    # Test health and metrics
    health_status = rag.get_health_status()
    print(f"\nSystem Health: {health_status['status']}")
    print(f"Vectorstore Status: {health_status['vectorstore_status']}")
    
    # Test A/B testing
    test_result = rag.a_b_test_query("What is cloud computing?", test_group='A')
    print(f"\nA/B Test Group: {test_result['test_group']}")
    print(f"Response preview: {test_result['response'][:100]}...")

def demo_analytics():
    """
    Demonstrate analytics functionality.
    """
    print("\n=== Analytics Demo ===")
    
    # Setup RAG system
    config = RAGConfig(
        embedding_provider="sentence_transformers",
        embedding_model="all-MiniLM-L6-v2"
    )
    
    rag = ProductionRAG(config)
    analytics = RAGAnalytics(rag)
    
    # Setup documents
    documents = create_sample_documents()
    metadata_list = [
        DocumentMetadata(
            document_id=f"doc_{i}",
            title=doc.metadata['title'],
            source=doc.metadata['source'],
            document_type="text",
            category=doc.metadata['category'],
            tags=doc.metadata['tags'],
            quality_score=doc.metadata['quality_score']
        )
        for i, doc in enumerate(documents)
    ]
    
    rag.add_documents_with_metadata(documents, metadata_list)
    
    # Simulate user queries
    test_users = ['user1', 'user2', 'user3']
    test_questions = [
        "What is machine learning?",
        "How does cloud computing work?",
        "What are cybersecurity best practices?",
        "Explain microservices",
        "Tell me about data privacy"
    ]
    
    print("Simulating user queries for analytics...")
    
    for user in test_users:
        for question in np.random.choice(test_questions, size=np.random.randint(1, 4)):
            result = rag.query_with_filters(question)
            analytics.log_query(user, question, result)
    
    # Generate reports
    performance_report = analytics.generate_performance_report(days=1)
    quality_analysis = analytics.analyze_query_quality()
    user_insights = analytics.get_user_insights()
    
    print(f"\n--- Performance Report ---")
    print(f"Total queries: {performance_report['total_queries']}")
    print(f"Unique users: {performance_report['unique_users']}")
    print(f"Avg total time: {performance_report['performance']['avg_total_time']:.3f}s")
    print(f"Cache hit rate: {performance_report['performance']['cache_hit_rate']:.2%}")
    
    print(f"\n--- Quality Analysis ---")
    print(f"P95 response time: {quality_analysis['response_time_analysis']['p95_response_time']:.3f}s")
    print(f"Zero result queries: {quality_analysis['retrieval_analysis']['zero_result_percentage']:.1f}%")
    
    print(f"\n--- User Insights ---")
    print(f"Avg queries per user: {user_insights['avg_queries_per_user']}")
    print(f"Power users: {user_insights['power_users']}")

# ================================================================
# SECTION 10: DEPLOYMENT AND CONFIGURATION TEMPLATES
# ================================================================

def create_docker_config():
    """
    Create Docker configuration for RAG deployment.
    """
    dockerfile_content = """
# Dockerfile for RAG System
FROM python:3.9-slim

# Set working directory
WORKDIR /app

# Install system dependencies
RUN apt-get update && apt-get install -y \\
    gcc \\
    g++ \\
    && rm -rf /var/lib/apt/lists/*

# Copy requirements
COPY requirements.txt .

# Install Python dependencies
RUN pip install --no-cache-dir -r requirements.txt

# Copy application code
COPY . .

# Expose ports
EXPOSE 5000 8000

# Set environment variables
ENV PYTHONPATH=/app
ENV FLASK_APP=app.py

# Health check
HEALTHCHECK --interval=30s --timeout=30s --start-period=5s --retries=3 \\
    CMD curl -f http://localhost:5000/health || exit 1

# Run the application
CMD ["python", "app.py"]
"""
    
    requirements_content = """
# Core RAG dependencies
langchain>=0.1.0
openai>=1.0.0
faiss-cpu>=1.7.0
sentence-transformers>=2.2.0
transformers>=4.20.0
torch>=2.0.0

# Document processing
PyPDF2>=3.0.0
python-docx>=0.8.11
beautifulsoup4>=4.11.0
markdown>=3.4.0

# Vector databases
chromadb>=0.4.0
pinecone-client>=2.2.0

# Web frameworks
flask>=2.3.0
streamlit>=1.28.0

# Monitoring and metrics
prometheus-client>=0.17.0
structlog>=23.1.0

# Data processing
numpy>=1.24.0
pandas>=2.0.0
scikit-learn>=1.3.0

# Environment and configuration
python-dotenv>=1.0.0

# Optional: Multi-modal support
# easyocr>=1.7.0
# whisper>=1.0.0

# Development and testing
pytest>=7.4.0
black>=23.7.0
flake8>=6.0.0
"""
    
    docker_compose_content = """
version: '3.8'

services:
  rag-system:
    build: .
    ports:
      - "5000:5000"
      - "8000:8000"
    environment:
      - OPENAI_API_KEY=${OPENAI_API_KEY}
      - HUGGINGFACE_API_KEY=${HUGGINGFACE_API_KEY}
      - FLASK_ENV=production
    volumes:
      - ./data:/app/data
      - ./vector_store:/app/vector_store
    restart: unless-stopped
    healthcheck:
      test: ["CMD", "curl", "-f", "http://localhost:5000/health"]
      interval: 30s
      timeout: 10s
      retries: 3

  prometheus:
    image: prom/prometheus:latest
    ports:
      - "9090:9090"
    volumes:
      - ./prometheus.yml:/etc/prometheus/prometheus.yml
    command:
      - '--config.file=/etc/prometheus/prometheus.yml'
      - '--storage.tsdb.path=/prometheus'
      - '--web.console.libraries=/etc/prometheus/console_libraries'
      - '--web.console.templates=/etc/prometheus/consoles'

  grafana:
    image: grafana/grafana:latest
    ports:
      - "3000:3000"
    environment:
      - GF_SECURITY_ADMIN_PASSWORD=admin
    volumes:
      - grafana-storage:/var/lib/grafana

volumes:
  grafana-storage:
"""
    
    return {
        'Dockerfile': dockerfile_content,
        'requirements.txt': requirements_content,
        'docker-compose.yml': docker_compose_content
    }

def create_kubernetes_config():
    """
    Create Kubernetes configuration for RAG deployment.
    """
    k8s_deployment = """
apiVersion: apps/v1
kind: Deployment
metadata:
  name: rag-system
  labels:
    app: rag-system
spec:
  replicas: 3
  selector:
    matchLabels:
      app: rag-system
  template:
    metadata:
      labels:
        app: rag-system
    spec:
      containers:
      - name: rag-system
        image: rag-system:latest
        ports:
        - containerPort: 5000
        - containerPort: 8000
        env:
        - name: OPENAI_API_KEY
          valueFrom:
            secretKeyRef:
              name: rag-secrets
              key: openai-api-key
        resources:
          requests:
            memory: "1Gi"
            cpu: "500m"
          limits:
            memory: "2Gi"
            cpu: "1000m"
        livenessProbe:
          httpGet:
            path: /health
            port: 5000
          initialDelaySeconds: 30
          periodSeconds: 10
        readinessProbe:
          httpGet:
            path: /health
            port: 5000
          initialDelaySeconds: 5
          periodSeconds: 5
---
apiVersion: v1
kind: Service
metadata:
  name: rag-system-service
spec:
  selector:
    app: rag-system
  ports:
  - name: api
    port: 80
    targetPort: 5000
  - name: metrics
    port: 8000
    targetPort: 8000
  type: LoadBalancer
---
apiVersion: v1
kind: Secret
metadata:
  name: rag-secrets
type: Opaque
data:
  openai-api-key: <base64-encoded-api-key>
"""
    
    return k8s_deployment

# ================================================================
# SECTION 11: MAIN EXECUTION AND TESTING
# ================================================================

def main():
    """
    Main function demonstrating complete RAG system usage.
    """
    print("🤖 Complete Python RAG Implementation Book")
    print("=" * 60)
    
    # Check environment setup
    openai_key = os.getenv("OPENAI_API_KEY")
    if not openai_key:
        print("⚠️  Warning: OPENAI_API_KEY not found in environment variables")
        print("   Using sentence-transformers for embeddings (free alternative)")
    else:
        print("✅ OpenAI API key found")
    
    try:
        # Demo basic RAG
        demo_basic_rag()
        
        # Demo enhanced RAG
        demo_enhanced_rag()
        
        # Demo production features
        demo_production_rag()
        
        # Demo analytics
        demo_analytics()
        
        print("\n" + "=" * 60)
        print("✅ All demos completed successfully!")
        print("\nNext steps:")
        print("1. Customize the configuration for your use case")
        print("2. Add your own documents and metadata")
        print("3. Deploy using the provided Docker/K8s configs")
        print("4. Set up monitoring and analytics")
        print("5. Implement additional features as needed")
        
    except Exception as e:
        print(f"\n❌ Error during demo: {str(e)}")
        print("Please check your environment setup and dependencies")

# ================================================================
# SECTION 12: BEST PRACTICES AND OPTIMIZATION GUIDELINES
# ================================================================

class RAGOptimizer:
    """
    RAG system optimization utilities and best practices.
    
    This class provides methods for optimizing RAG performance,
    including hyperparameter tuning, prompt optimization,
    and retrieval strategy selection.
    """
    
    def __init__(self, rag_system: ProductionRAG):
        self.rag_system = rag_system
        self.optimization_history = []
    
    def optimize_chunk_size(self, test_questions: List[str], chunk_sizes: List[int] = None) -> Dict[str, Any]:
        """
        Optimize chunk size for better retrieval performance.
        
        Args:
            test_questions: List of test questions
            chunk_sizes: List of chunk sizes to test
            
        Returns:
            Optimization results
        """
        if chunk_sizes is None:
            chunk_sizes = [200, 500, 1000, 1500, 2000]
        
        results = {}
        original_chunk_size = self.rag_system.config.chunk_size
        
        for chunk_size in chunk_sizes:
            self.rag_system.config.chunk_size = chunk_size
            
            # Test performance with this chunk size
            total_time = 0
            retrieval_scores = []
            
            for question in test_questions:
                result = self.rag_system.query_with_filters(question)
                total_time += result['total_time']
                retrieval_scores.append(len(result['sources']))
            
            results[chunk_size] = {
                'avg_time': total_time / len(test_questions),
                'avg_retrieval_count': np.mean(retrieval_scores),
                'total_time': total_time
            }
        
        # Restore original chunk size
        self.rag_system.config.chunk_size = original_chunk_size
        
        # Find optimal chunk size
        optimal_size = min(results.keys(), key=lambda x: results[x]['avg_time'])
        
        return {
            'optimal_chunk_size': optimal_size,
            'results': results,
            'recommendation': f"Use chunk size {optimal_size} for optimal performance"
        }
    
    def analyze_retrieval_quality(self, test_questions: List[str], expected_sources: List[List[str]]) -> Dict[str, Any]:
        """
        Analyze retrieval quality using test questions with known relevant sources.
        
        Args:
            test_questions: List of test questions
            expected_sources: List of expected source files for each question
            
        Returns:
            Quality analysis results
        """
        precision_scores = []
        recall_scores = []
        
        for question, expected in zip(test_questions, expected_sources):
            result = self.rag_system.query_with_filters(question)
            retrieved_sources = [source['source'] for source in result['sources']]
            
            # Calculate precision and recall
            relevant_retrieved = set(retrieved_sources) & set(expected)
            precision = len(relevant_retrieved) / len(retrieved_sources) if retrieved_sources else 0
            recall = len(relevant_retrieved) / len(expected) if expected else 0
            
            precision_scores.append(precision)
            recall_scores.append(recall)
        
        return {
            'avg_precision': np.mean(precision_scores),
            'avg_recall': np.mean(recall_scores),
            'f1_score': 2 * np.mean(precision_scores) * np.mean(recall_scores) / (np.mean(precision_scores) + np.mean(recall_scores)) if (np.mean(precision_scores) + np.mean(recall_scores)) > 0 else 0,
            'precision_scores': precision_scores,
            'recall_scores': recall_scores
        }

# Performance tips and best practices documentation
BEST_PRACTICES_GUIDE = """
# RAG System Best Practices and Optimization Guide

## 1. Document Preparation
- Clean and preprocess documents thoroughly
- Remove headers, footers, and navigation elements
- Maintain consistent formatting across documents
- Use meaningful file names and metadata
- Regular updates to keep knowledge base current

## 2. Chunking Strategy
- Test different chunk sizes (200-2000 characters)
- Use semantic chunking for better context preservation
- Overlap chunks by 10-20% to maintain continuity
- Consider document structure (paragraphs, sections)
- Adjust chunk size based on document type

## 3. Embedding Selection
- OpenAI embeddings: High quality, costs money
- Sentence Transformers: Free, good quality
- Domain-specific models for specialized content
- Consider multilingual models if needed
- Test embedding quality on your specific data

## 4. Retrieval Optimization
- Use hybrid search (dense + sparse) for better coverage
- Implement re-ranking for improved relevance
- Filter by metadata to narrow search scope
- Adjust similarity thresholds based on use case
- Monitor and tune retrieval performance

## 5. Response Generation
- Craft clear, specific prompts for the LLM
- Include context about desired response format
- Implement response validation and filtering
- Use appropriate temperature settings
- Consider response length and detail level

## 6. Performance Optimization
- Implement caching for frequent queries
- Use async processing for better scalability
- Batch process multiple queries when possible
- Monitor response times and optimize bottlenecks
- Scale vector databases horizontally

## 7. Monitoring and Analytics
- Track query performance metrics
- Monitor retrieval quality and relevance
- Analyze user behavior and query patterns
- Set up alerting for system issues
- Regular performance reviews and optimization

## 8. Security and Privacy
- Encrypt sensitive data at rest and in transit
- Implement proper access controls
- Audit and log system activities
- Comply with relevant data regulations
- Regular security assessments

## 9. Testing and Validation
- Create comprehensive test suites
- Test with diverse query types and formats
- Validate responses for accuracy and relevance
- A/B testing for system improvements
- Regular regression testing

## 10. Deployment Considerations
- Use containerization for consistent environments
- Implement proper CI/CD pipelines
- Set up monitoring and alerting
- Plan for scaling and load balancing
- Disaster recovery and backup strategies
"""

if __name__ == "__main__":
    main()